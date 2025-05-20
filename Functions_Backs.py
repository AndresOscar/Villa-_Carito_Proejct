import pandas as pd
from docx import Document
from docx.oxml.ns import qn
import numpy as np
from copy import deepcopy
from openpyxl import load_workbook
from openpyxl.utils import coordinate_to_tuple, get_column_letter
from functios_database import inicializar_base_datos,obtener_configuraciones,guardar_configuracion

def detectar_rango_tabla(excel_path, celda_inicial="B2", sheet_name=None):
    """
    Detecta el rango de una tabla en Excel comenzando desde la celda inicial.
    Devuelve el rango como string, por ejemplo 'B2:F16'.
    """
    wb = load_workbook(excel_path, data_only=True)
    ws = wb[sheet_name] if sheet_name else wb.active

    start_col_letter, start_row = coordinate_to_tuple(celda_inicial)
    start_col_index = start_col_letter  # en realidad devuelve número, no letra
    start_col_letter = get_column_letter(start_col_index)

    # Convertir a índice 0-based
    row = start_row
    col = start_col_index

    # Detectar límite de columnas
    last_col = col
    while ws.cell(row=row, column=last_col).value not in (None, ""):
        last_col += 1
    last_col -= 1  # retroceder al último con valor

    # Detectar límite de filas
    last_row = row
    while ws.cell(row=last_row, column=col).value not in (None, ""):
        last_row += 1
    last_row -= 1  # retroceder al último con valor

    col_end_letter = get_column_letter(last_col)

    rango = f"{start_col_letter}{start_row}:{col_end_letter}{last_row}"
    return rango


def read_excel_table(excel_path, sheet_name, excel_range):
    """Lee un rango de una hoja de Excel y lo devuelve como DataFrame."""
    import re
    match = re.match(r'([A-Z]+)(\d+):([A-Z]+)(\d+)', excel_range)
    if not match:
        raise ValueError("Formato de rango inválido. Usa algo como 'B2:L14'.")
    col_start, row_start, col_end, row_end = match.groups()
    row_start = int(row_start)
    row_end = int(row_end)
    cols = f"{col_start}:{col_end}"

    # Leer solo columnas especificadas
    df = pd.read_excel(excel_path, sheet_name=sheet_name, usecols=cols)

    # Seleccionar solo filas dentro del rango
    df = df.iloc[row_start - 2 : row_end - 1].reset_index(drop=True)
    return df

def find_paragraph_with_label(doc, label):
    """Encuentra el párrafo que contiene la etiqueta específica."""
    for paragraph in doc.paragraphs:
        if label in paragraph.text:
            return paragraph
    return None

def find_table_after_paragraph(doc, paragraph):
    """Encuentra la tabla inmediatamente después del párrafo dado."""
    para_element = paragraph._p
    for tbl in doc.tables:
        if tbl._element.getprevious() == para_element:
            return tbl
    # También podría ser la siguiente en el flujo de contenido
    for tbl in doc.tables:
        if para_element in tbl._element.xpath('.//preceding::w:p'):
            return tbl
    return None


def get_grid_span(cell):
    """Devuelve cuántas columnas abarca la celda (por gridSpan), por defecto 1."""
    tc = cell._tc
    grid_span = tc.xpath('.//w:gridSpan')
    if grid_span:
        return int(grid_span[0].get(qn('w:val')))
    return 1

def detectar_filas_con_columnas_unidas(table):
    """
    Devuelve una lista con los índices de las filas que contienen
    al menos una celda fusionada horizontalmente (gridSpan > 1).
    """
    filas_con_fusiones = []
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if get_grid_span(cell) > 1:
                filas_con_fusiones.append(i)
                break  # solo necesitamos marcar la fila una vez
    return filas_con_fusiones


def update_table_cells(word_table, df):
    """Actualiza las celdas de una tabla de Word con los valores de un DataFrame,
    omitiendo las filas fusionadas (no las modifica), pero manteniendo el orden."""
    
    filas_con_fusiones = detectar_filas_con_columnas_unidas(word_table)
    print(f"Filas fusionadas detectadas (se omiten): {filas_con_fusiones}")

    if len(df) > len(word_table.rows):
        print(f"Advertencia: Excel tiene {len(df)} filas, pero la tabla de Word solo tiene {len(word_table.rows)} filas.")

    for i in range(len(word_table.rows)):
        if i >= len(df):
            break  # no hay más filas de Excel para insertar

        if i in filas_con_fusiones:
            print(f"Fila {i} fusionada: se omite actualización pero se avanza índice de Excel.")
            continue  # NO escribimos en filas fusionadas, pero igual avanzamos Excel

        row = df.iloc[i]  # ✅ usar el mismo índice i para Excel y Word
        for j, value in enumerate(row):
            if j >= len(word_table.columns):
                print(f"Advertencia: columna {j} de Excel excede columnas de la tabla Word.")
                break
            word_cell = word_table.cell(i, j)
            word_cell.text = "" if pd.isna(value) else str(value)




def update_word_table_from_excel(excel_path, sheet_name, excel_range, word_path, label, output_path=None):
   
    """Actualiza una tabla existente en Word con los datos de una tabla de Excel."""
    print(excel_range)

    df = read_excel_table(excel_path, sheet_name, excel_range)
    doc = Document(word_path)

    paragraph = find_paragraph_with_label(doc, label)
    if not paragraph:
        print(f"No se encontró la etiqueta '{label}' en el documento.")
        return

    table = find_table_after_paragraph(doc, paragraph)

    if not table:
        print("No se encontró una tabla después de la etiqueta.")
        return
    
    
    num_columnas_excel = df.shape[1]
    num_columnas_word = len(table.columns)

    if num_columnas_excel != num_columnas_word:
        print(f"Error: La tabla de Excel tiene {num_columnas_excel} columnas, pero la tabla de Word tiene {num_columnas_word} columnas.")
        print("No se puede actualizar la tabla porque las columnas no coinciden.")
        return


    ajustar_tabla_word(table, df.shape[0], df.shape[1])


    update_table_cells(table, df)
    print("Tabla actualizada con los nuevos valores.")

    save_path = output_path if output_path else word_path
    doc.save(save_path)
    print(f"Documento guardado en {save_path}")




def ajustar_tabla_word(table, num_filas_excel, num_columnas_excel):
    num_filas_word = len(table.rows)
    num_columnas_word = len(table.columns)

    print(f"Tamaño actual Word: {num_filas_word} filas x {num_columnas_word} columnas")
    print(f"Tamaño requerido Excel: {num_filas_excel} filas x {num_columnas_excel} columnas")

    # Ajustar filas
    if num_filas_word < num_filas_excel:
        for _ in range(num_filas_excel - num_filas_word):
            table.add_row()
        print(f"Se agregaron {num_filas_excel - num_filas_word} filas.")
    elif num_filas_word > num_filas_excel:
        for _ in range(num_filas_word - num_filas_excel):
            table._tbl.remove(table.rows[-1]._tr)
        print(f"Se eliminaron {num_filas_word - num_filas_excel} filas.")

    # Ajustar columnas
    if num_columnas_word < num_columnas_excel:
        for row in table.rows:
            for _ in range(num_columnas_excel - num_columnas_word):
                row._tr.append(deepcopy(row.cells[-1]._tc))
        print(f"Se agregaron {num_columnas_excel - num_columnas_word} columnas.")
    elif num_columnas_word > num_columnas_excel:
        for row in table.rows:
            for _ in range(num_columnas_word - num_columnas_excel):
                row._tr.remove(row.cells[-1]._tc)
        print(f"Se eliminaron {num_columnas_word - num_columnas_excel} columnas.")


def actualizar_todas_las_tablas():
    configs = obtener_configuraciones()
    for config in configs:
        id_, excel_file, sheet_name, excel_range, word_file, table_label, output_file = config
        try:
            print(f"Actualizando tabla {id_}...")
            update_word_table_from_excel(excel_file, sheet_name, excel_range, word_file, table_label, output_file)
            print(f"Tabla {id_} actualizada correctamente.")
        except Exception as e:
            print(f"Error al actualizar tabla {id_}: {e}")

from docx.shared import Pt

def formatear_tabla(tabla, fuente='Arial', tamaño=10):
    """
    Aplica formato de fuente y tamaño a todas las celdas de una tabla de Word.
    También aplica bordes a toda la tabla.
    """
    for fila in tabla.rows:
        for celda in fila.cells:
            for parrafo in celda.paragraphs:
                for run in parrafo.runs:
                    run.font.name = fuente
                    run.font.size = Pt(tamaño)

    aplicar_bordes_a_tabla(tabla)

def aplicar_bordes_a_tabla(tabla):
    tbl = tabla._tbl  # accedemos al XML de la tabla

    tblBorders = tbl.xpath('.//w:tblBorders')
    if tblBorders:
        tblBorders = tblBorders[0]
    else:
        # Crear nodo de bordes si no existe
        from lxml import etree
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        tblPr = tbl.xpath('./w:tblPr', namespaces=namespaces)[0]
        tblBorders = etree.SubElement(tblPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')

    for borde in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        tag = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}' + borde
        elemento = tblBorders.find(tag)
        if elemento is None:
            from lxml import etree
            elemento = etree.SubElement(tblBorders, tag)
        elemento.set('val', 'single')
        elemento.set('sz', '4')      # tamaño del borde
        elemento.set('space', '0')
        elemento.set('color', '000000')  # negro


def obtener_formato_tabla(tabla):
    """
    Retorna una tupla (fuente, tamaño) del primer texto encontrado en la tabla.
    Si no encuentra fuente o tamaño, devuelve (None, None).
    """
    for fila in tabla.rows:
        for celda in fila.cells:
            for parrafo in celda.paragraphs:
                for run in parrafo.runs:
                    fuente = run.font.name
                    tamaño = run.font.size.pt if run.font.size else None
                    if fuente or tamaño:
                        return (fuente, tamaño)
    return (None, None)  # si no encuentra nada



def find_table_by_label(doc, label_text):
    for i, paragraph in enumerate(doc.paragraphs):
        if label_text in paragraph.text:
            for element in doc.element.body.iter():
                if element.tag.endswith('p') and element == paragraph._element:
                    next_element = element.getnext()
                    while next_element is not None:
                        if next_element.tag.endswith('tbl'):
                            for tbl in doc.tables:
                                if tbl._element == next_element:
                                    return tbl
                        next_element = next_element.getnext()
    return None

def format_table_money_columns(table, money_cols, header_rows=1):
    for row in table.rows[header_rows:]:
        for col_idx in money_cols:
            cell_text = row.cells[col_idx].text.strip()
            try:
                monto = float(cell_text.replace(",", "").replace("$", ""))
                row.cells[col_idx].text = "${:,.2f}".format(monto)
            except ValueError:
                pass  # dejar sin cambiar si falla

def main(doc_name, label, money_columns, header_rows):
    # Cargar documento
    doc = Document(doc_name)

    # Buscar tabla por etiqueta
    table = find_table_by_label(doc, label)

    if table:
        format_table_money_columns(table, money_columns, header_rows)
        doc.save(doc_name)  # sobrescribe el mismo archivo
        print(f"✅ Tabla '{label}' encontrada y columnas de dinero formateadas.")
    else:
        print(f"⚠ No se encontró la tabla con la etiqueta '{label}'.")


# ===========================
# Ejemplo de uso
# ===========================
def main_1():

    excel_file   = r"C:\Users\aguti\Desktop\Comunicaciones Villa Craito\Atutomaizacion_Informes\BASE DE DATOS BENEFICIARIOS VILLA CARITO.xlsx"
    sheet_name   = "CONSOLIDADO ESTADO BENEFICIARIO"
    excel_range  = "B2:L14"
    word_file    = r"C:\Users\aguti\Desktop\Comunicaciones Villa Craito\Atutomaizacion_Informes\023 (VER02) INFORME GENERAL 13 DE MARZO DE 2025.docx"
    table_label  = "2. CONSOLIDADO DE ESTADO DE BENEFICIARIOS."
    #output_file  = r"C:\Users\aguti\Desktop\Comunicaciones Villa Craito\Atutomaizacion_Informes\Doc1_updated.docx"

    update_word_table_from_excel(excel_file, sheet_name, excel_range, word_file, table_label)

