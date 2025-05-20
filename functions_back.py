from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

def col_to_letter(n):
    """Convierte un número de columna a su letra equivalente (1 -> 'A')."""
    string = ""
    while n > 0:
        n, remainder = divmod(n - 1, 26)
        string = chr(65 + remainder) + string
    return string

def hex_to_rgb(hex_color):
    """Convierte color HEX a RGBColor"""
    hex_color = hex_color.lstrip('#')
    return RGBColor(*tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4)))

def get_default_format():
    """Devuelve un formato predeterminado para celdas vacías"""
    return {
        'value': '',
        'font': {
            'bold': False,
            'size': 11,
            'color': '#000000',
            'name': 'Calibri'
        },
        'background': '#FFFFFF',
        'borders': {
            'top': {'style': 'NONE', 'color': '#000000'},
            'bottom': {'style': 'NONE', 'color': '#000000'},
            'left': {'style': 'NONE', 'color': '#000000'},
            'right': {'style': 'NONE', 'color': '#000000'}
        },
        'alignment': {
            'horizontal': 'LEFT',
            'vertical': 'TOP'
        },
        'text_format': {
            'wrap': False,
            'rotation': 0
        }
    }


def get_sheet_data2(sheets_service, sheet_id, sheet_name, start_row, end_row, start_col, end_col):
    """
    Extrae datos y TODOS los formatos de un rango específico de Google Sheets.
    Devuelve una tabla (lista de filas) y el número de columnas.
    """
    start_col_letter = col_to_letter(start_col)
    end_col_letter = col_to_letter(end_col)
    range_notation = f"{sheet_name}!{start_col_letter}{start_row}:{end_col_letter}{end_row}"
    
    try:
        sheet_response = sheets_service.spreadsheets().get(
            spreadsheetId=sheet_id,
            ranges=[range_notation],
            includeGridData=True,
            fields="sheets(data(rowData(values(formattedValue,effectiveFormat,textFormatRuns)))"
        ).execute()

        grid_data = sheet_response['sheets'][0]['data'][0].get('rowData', [])
        num_cols = end_col - start_col + 1
        tabla = []

        for row in grid_data:
            fila = []
            values = row.get('values', [])
            
            for i in range(num_cols):
                if i < len(values):
                    cell = values[i]
                    valor = cell.get('formattedValue', '')
                    fmt = cell.get('effectiveFormat', {})
                    
                    # Extracción de formato de texto
                    text_format = fmt.get('textFormat', {})
                    text_runs = cell.get('textFormatRuns', [])
                    
                    # Extracción de formato de celda
                    background = fmt.get('backgroundColor', {}).get('rgbColor', {}).get('red', 1)
                    bg_color = f"#{int(background.get('red', 1)*255):02x}{int(background.get('green', 1)*255):02x}{int(background.get('blue', 1)*255):02x}"
                    
                    # Manejo de bordes (simplificado)
                    borders = {}
                    for side in ['top', 'bottom', 'left', 'right']:
                        border = fmt.get('borders', {}).get(side, {})
                        borders[side] = {
                            'style': border.get('style', 'NONE'),
                            'color': f"#{int(border.get('color', {}).get('red', 0)*255):02x}"
                                          f"{int(border.get('color', {}).get('green', 0)*255):02x}"
                                          f"{int(border.get('color', {}).get('blue', 0)*255):02x}"
                        }
                    
                    fila.append({
                        'value': valor,
                        'font': {
                            'bold': text_format.get('bold', False),
                            'size': text_format.get('fontSize', 11),
                            'color': f"#{int(text_format.get('foregroundColor', {}).get('red', 0)*255):02x}"
                                      f"{int(text_format.get('foregroundColor', {}).get('green', 0)*255):02x}"
                                      f"{int(text_format.get('foregroundColor', {}).get('blue', 0)*255):02x}",
                            'name': text_format.get('fontFamily', 'Calibri')
                        },
                        'background': bg_color,
                        'borders': borders,
                        'alignment': {
                            'horizontal': fmt.get('horizontalAlignment', 'LEFT').upper(),
                            'vertical': fmt.get('verticalAlignment', 'TOP').upper()
                        },
                        'text_format': {
                            'wrap': fmt.get('wrapStrategy') == 'WRAP',
                            'rotation': fmt.get('textRotation', {}).get('angle', 0)
                        }
                    })
                else:
                    fila.append(get_default_format())
                    
            tabla.append(fila)
            
        return tabla, num_cols
        
    except Exception as e:
        print(f"Error al obtener datos de Sheets: {str(e)}")
        # Devuelve tabla vacía con formato predeterminado
        return [[get_default_format() for _ in range(end_col - start_col + 1)] 
                for _ in range(end_row - start_row + 1)], end_col - start_col + 1
    

def create_table22(table, num_cols, document):
    """
    Crea una tabla en Word con todos los formatos extraídos de Google Sheets
    """
    try:
        # Crear tabla básica
        doc_table = document.add_table(rows=len(table), cols=num_cols)
        doc_table.style = 'Table Grid'
        doc_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Configuración de layout fijo
        tbl = doc_table._tbl
        tblPr = tbl.tblPr
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        
        # Rellenar tabla con datos y formato
        for i, fila in enumerate(table):
            for j, cell_data in enumerate(fila):
                cell = doc_table.cell(i, j)
                cell.text = ""
                
                # Aplicar formato de celda
                apply_cell_formatting(cell, cell_data)
                
                # Añadir texto con formato
                p = cell.paragraphs[0]
                run = p.add_run(cell_data['value'])
                
                # Aplicar formato de texto
                run.font.bold = cell_data['font']['bold']
                run.font.size = Pt(cell_data['font']['size'])
                run.font.color.rgb = hex_to_rgb(cell_data['font']['color'])
                run.font.name = cell_data['font']['name']
                
                # Aplicar alineación
                p.alignment = map_alineacion(cell_data['alignment']['horizontal'])
                
                # Aplicar rotación de texto (si está soportada)
                if cell_data['text_format']['rotation'] not in [0, 180, 270]:
                    text_direction = OxmlElement('w:textDirection')
                    text_direction.set(qn('w:val'), 'btLr')  # Bottom-to-top, left-to-right
                    p._p.get_or_add_pPr().append(text_direction)
        
        return doc_table
        
    except Exception as e:
        print(f"Error al crear tabla en Word: {str(e)}")
        raise



def apply_cell_formatting(cell, cell_data):
    """
    Aplica formato avanzado a una celda de Word
    """
    # Fondo de celda
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), cell_data['background'][1:])  # Elimina el #
    cell._tc.get_or_add_tcPr().append(shading)
    
    # Bordes (implementación simplificada)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), cell_data['borders'][side]['style'].lower())
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), cell_data['borders'][side]['color'][1:])
        tcBorders.append(border)
    cell._tc.get_or_add_tcPr().append(tcBorders)
    
    # Ajuste de texto
    tcPr = cell._tc.get_or_add_tcPr()
    noWrap = OxmlElement('w:noWrap')
    if cell_data['text_format']['wrap']:
        noWrap.set(qn('w:val'), '0')
    else:
        noWrap.set(qn('w:val'), '1')
    tcPr.append(noWrap)
    
    # Alineación vertical
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), cell_data['alignment']['vertical'].lower())
    tcPr.append(vAlign)


def map_alineacion(gsheets_alignment):
    """
    Mapea alineación de Google Sheets a Word
    """
    mapping = {
        'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
        'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
        'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
        'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    return mapping.get(gsheets_alignment, WD_ALIGN_PARAGRAPH.LEFT)
